from __future__ import annotations

import copy
import hashlib
import html
import io
import re
from pathlib import Path
from typing import Any

from docx import Document

from app.services.jd_parser import JDParserService
from app.services.job_targeted_scorer import JobProfileData, JobTargetedScorer
from app.services.local_reasoner import LocalReasoner
from app.services.resume_parser import ResumeParserService
from app.services.skill_taxonomy import canonicalize_skill, expand_skill_aliases

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}")
_URL_RE = re.compile(r"https?://[^\s]+|www\.[^\s]+")
_BULLET_RE = re.compile(r"^\s*(?:[-*•●◦▪]|\d+[.)])\s*")
_SECTION_KEYS: dict[str, str] = {
    "summary": "summary",
    "profile": "summary",
    "objective": "summary",
    "skills": "skills",
    "core competencies": "skills",
    "technical skills": "skills",
    "experience": "experience",
    "work experience": "experience",
    "employment": "experience",
    "projects": "projects",
    "education": "education",
    "certifications": "certifications",
    "awards": "awards",
}

_DEFAULT_TEMPLATE_SETTINGS = {
    "font_scale": "normal",
    "density": "normal",
    "accent": "slate",
    "show_icons": False,
    "show_sections": {
        "summary": True,
        "skills": True,
        "experience": True,
        "projects": True,
        "education": True,
        "certifications": True,
        "awards": True,
    },
}


def _dedupe(items: list[str]) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for item in items:
        normalized = item.strip().lower()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        output.append(item.strip())
    return output


def _clean_line(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _section_from_heading(line: str) -> str | None:
    clean = line.strip().lower().rstrip(":")
    for key, mapped in _SECTION_KEYS.items():
        if clean == key:
            return mapped
    return None


def default_structured_resume(title: str = "Untitled Resume") -> dict[str, Any]:
    return {
        "header": {
            "name": title if title != "Untitled Resume" else "",
            "title": "",
            "email": "",
            "phone": "",
            "location": "",
            "links": [],
        },
        "summary": "",
        "skills": {"categories": []},
        "experience": {"items": []},
        "projects": {"items": []},
        "education": {"items": []},
        "certifications": [],
        "awards": [],
        "ats_keywords": [],
        "evidence_map": {},
    }


class ResumeStudioService:
    def __init__(self) -> None:
        self.resume_parser = ResumeParserService()
        self.jd_parser = JDParserService()
        self.scorer = JobTargetedScorer()
        self.reasoner = LocalReasoner()

    async def parse_upload(self, filename: str, content: bytes) -> str:
        return await self.resume_parser.parse(filename, content)

    def ensure_schema(self, structured: dict[str, Any] | None) -> dict[str, Any]:
        base = default_structured_resume()
        if structured is None:
            return base
        merged = copy.deepcopy(base)
        for key, value in structured.items():
            if key in {"header", "skills", "experience", "projects", "education"} and isinstance(value, dict):
                merged[key] = {**merged[key], **value}
            else:
                merged[key] = value
        if not isinstance(merged.get("ats_keywords"), list):
            merged["ats_keywords"] = []
        if not isinstance(merged.get("evidence_map"), dict):
            merged["evidence_map"] = {}
        return merged

    def parse_resume_text(self, text: str, title_hint: str = "Untitled Resume") -> dict[str, Any]:
        lines = [_clean_line(line) for line in text.splitlines()]
        lines = [line for line in lines if line]

        structured = default_structured_resume(title_hint)
        header_lines: list[str] = []
        sections: dict[str, list[str]] = {
            "summary": [],
            "skills": [],
            "experience": [],
            "projects": [],
            "education": [],
            "certifications": [],
            "awards": [],
        }
        current_section = "summary"

        for index, line in enumerate(lines):
            maybe_section = _section_from_heading(line)
            if maybe_section:
                current_section = maybe_section
                continue
            if index < 6 and current_section == "summary":
                header_lines.append(line)
                continue
            sections[current_section].append(line)

        structured["header"] = self._parse_header(header_lines, title_hint)
        summary_lines = [
            line
            for line in sections["summary"]
            if not _EMAIL_RE.search(line) and not _PHONE_RE.search(line) and not _URL_RE.search(line)
        ]
        structured["summary"] = " ".join(summary_lines[:3]).strip()
        structured["skills"] = {"categories": self._parse_skills(sections["skills"])}
        structured["experience"] = {"items": self._parse_experience(sections["experience"])}
        structured["projects"] = {"items": self._parse_projects(sections["projects"])}
        structured["education"] = {"items": self._parse_education(sections["education"])}
        structured["certifications"] = _dedupe(sections["certifications"])
        structured["awards"] = _dedupe(sections["awards"])
        structured["ats_keywords"] = self._compute_ats_keywords(structured)
        structured["evidence_map"] = self._build_evidence_map(structured)
        return structured

    def structured_to_plain_text(self, structured_resume: dict[str, Any]) -> str:
        structured = self.ensure_schema(structured_resume)
        lines: list[str] = []
        header = structured["header"]
        if header.get("name"):
            lines.append(header["name"])
        if header.get("title"):
            lines.append(header["title"])
        for key in ("email", "phone", "location"):
            if header.get(key):
                lines.append(str(header[key]))
        for link in header.get("links", []):
            label = link.get("label", "link")
            url = link.get("url", "")
            if url:
                lines.append(f"{label}: {url}")

        if structured.get("summary"):
            lines.append("Summary")
            lines.append(structured["summary"])

        skill_categories = structured["skills"].get("categories", [])
        if skill_categories:
            lines.append("Skills")
            for category in skill_categories:
                lines.append(f"{category.get('name', 'Core')}: {', '.join(category.get('items', []))}")

        for item in structured["experience"].get("items", []):
            title = " | ".join(
                [entry for entry in [item.get("role"), item.get("company"), self._format_dates(item)] if entry]
            )
            if title:
                lines.append(title)
            lines.extend(item.get("bullets", []))

        for item in structured["projects"].get("items", []):
            if item.get("name"):
                lines.append(item["name"])
            lines.extend(item.get("bullets", []))

        for item in structured["education"].get("items", []):
            title = " | ".join(
                [entry for entry in [item.get("school"), item.get("degree"), self._format_dates(item)] if entry]
            )
            if title:
                lines.append(title)
            lines.extend(item.get("notes", []))

        lines.extend(structured.get("certifications", []))
        lines.extend(structured.get("awards", []))
        return "\n".join([line for line in lines if line]).strip()

    def render_resume_html(
        self,
        structured_resume: dict[str, Any],
        template_name: str = "ats_classic",
        template_settings: dict[str, Any] | None = None,
    ) -> str:
        structured = self.ensure_schema(structured_resume)
        settings = copy.deepcopy(_DEFAULT_TEMPLATE_SETTINGS)
        if template_settings:
            settings.update(template_settings)
            if "show_sections" in template_settings and isinstance(template_settings["show_sections"], dict):
                settings["show_sections"] = {
                    **_DEFAULT_TEMPLATE_SETTINGS["show_sections"],
                    **template_settings["show_sections"],
                }

        font_size = {"small": "14px", "normal": "15px", "large": "16px"}.get(
            settings.get("font_scale", "normal"),
            "15px",
        )
        spacing = {"compact": "0.42rem", "normal": "0.65rem", "comfortable": "0.82rem"}.get(
            settings.get("density", "normal"),
            "0.65rem",
        )

        accent_palette = {
            "slate": "#334155",
            "indigo": "#4338ca",
            "emerald": "#047857",
        }
        accent = accent_palette.get(str(settings.get("accent", "slate")).lower(), "#334155")

        normalized_template = template_name.strip().lower()
        if normalized_template == "modern_clean":
            section_border = "1px solid #e5e7eb"
            section_padding = "0.5rem 0 0.1rem"
        elif normalized_template in {"compact", "compact_onepager"}:
            section_border = "1px solid #f1f5f9"
            section_padding = "0.22rem 0 0.1rem"
        else:
            section_border = "1px solid #e2e8f0"
            section_padding = "0.45rem 0 0.1rem"

        header = structured["header"]
        links_html = " · ".join(
            [f"<a href='{html.escape(link.get('url', ''))}'>{html.escape(link.get('label', 'link'))}</a>" for link in header.get("links", []) if link.get("url")]
        )

        def section(title: str, body: str, key: str) -> str:
            if not settings["show_sections"].get(key, True) or not body.strip():
                return ""
            icon_map = {
                "summary": "✦",
                "skills": "◆",
                "experience": "▣",
                "projects": "◉",
                "education": "▤",
                "certifications": "✓",
                "awards": "★",
            }
            prefix = f"{icon_map.get(key, '')} " if settings.get("show_icons") else ""
            return (
                f"<section style='border-top:{section_border};padding:{section_padding};margin-top:{spacing};'>"
                f"<h2 style='margin:0 0 0.4rem;font-size:0.86rem;letter-spacing:.08em;text-transform:uppercase;color:{accent};'>{html.escape(prefix + title)}</h2>"
                f"{body}</section>"
            )

        summary_html = f"<p style='margin:0;line-height:1.55;'>{html.escape(structured.get('summary', ''))}</p>"
        skills_html = "".join(
            [
                "<div style='margin-bottom:.35rem;'><strong>"
                + html.escape(category.get("name", "Core"))
                + ":</strong> "
                + html.escape(", ".join(category.get("items", [])))
                + "</div>"
                for category in structured["skills"].get("categories", [])
            ]
        )
        experience_html = "".join(
            [
                "<article style='margin-bottom:.55rem;'>"
                + (
                    f"<p style='margin:0;font-weight:600;'>{html.escape(item.get('role', ''))}"
                    + (f" · {html.escape(item.get('company', ''))}" if item.get("company") else "")
                    + "</p>"
                )
                + (
                    f"<p style='margin:0;color:#475569;font-size:.9em;'>{html.escape(self._format_dates(item))}</p>"
                    if self._format_dates(item)
                    else ""
                )
                + "<ul style='margin:.3rem 0 0 1rem;padding:0;'>"
                + "".join([f"<li style='margin:.16rem 0;'>{html.escape(bullet)}</li>" for bullet in item.get("bullets", [])])
                + "</ul></article>"
                for item in structured["experience"].get("items", [])
            ]
        )
        projects_html = "".join(
            [
                "<article style='margin-bottom:.5rem;'>"
                + f"<p style='margin:0;font-weight:600;'>{html.escape(item.get('name', ''))}</p>"
                + "<ul style='margin:.3rem 0 0 1rem;padding:0;'>"
                + "".join([f"<li style='margin:.16rem 0;'>{html.escape(bullet)}</li>" for bullet in item.get("bullets", [])])
                + "</ul></article>"
                for item in structured["projects"].get("items", [])
            ]
        )
        education_html = "".join(
            [
                "<article style='margin-bottom:.45rem;'>"
                + f"<p style='margin:0;font-weight:600;'>{html.escape(item.get('school', ''))}</p>"
                + (
                    f"<p style='margin:0;color:#475569'>{html.escape(item.get('degree', ''))}</p>"
                    if item.get("degree")
                    else ""
                )
                + "</article>"
                for item in structured["education"].get("items", [])
            ]
        )
        certifications_html = "<ul style='margin:0 0 0 1rem;padding:0;'>" + "".join(
            [f"<li>{html.escape(item)}</li>" for item in structured.get("certifications", [])]
        ) + "</ul>"
        awards_html = "<ul style='margin:0 0 0 1rem;padding:0;'>" + "".join(
            [f"<li>{html.escape(item)}</li>" for item in structured.get("awards", [])]
        ) + "</ul>"

        return (
            "<!doctype html><html><head><meta charset='utf-8'/>"
            "<meta name='viewport' content='width=device-width, initial-scale=1'/>"
            "<title>Resume</title></head>"
            f"<body style='margin:0;background:#f8fafc;font-family:Inter,Segoe UI,Arial,sans-serif;color:#0f172a;font-size:{font_size};'>"
            "<main style='max-width:850px;margin:0 auto;background:white;padding:2rem 2rem 1.5rem;'>"
            f"<h1 style='margin:0;font-size:1.7rem;'>{html.escape(header.get('name', ''))}</h1>"
            + (f"<p style='margin:.35rem 0 .1rem;color:#334155;font-weight:600;'>{html.escape(header.get('title', ''))}</p>" if header.get("title") else "")
            + "<p style='margin:0;color:#475569;'>"
            + " · ".join([html.escape(v) for v in [header.get("email"), header.get("phone"), header.get("location")] if v])
            + (" · " if links_html and any([header.get("email"), header.get("phone"), header.get("location")]) else "")
            + links_html
            + "</p>"
            + section("Summary", summary_html, "summary")
            + section("Skills", skills_html, "skills")
            + section("Experience", experience_html, "experience")
            + section("Projects", projects_html, "projects")
            + section("Education", education_html, "education")
            + section("Certifications", certifications_html, "certifications")
            + section("Awards", awards_html, "awards")
            + "</main></body></html>"
        )

    def compute_structured_diff(
        self,
        base_structured: dict[str, Any],
        current_structured: dict[str, Any],
    ) -> dict[str, Any]:
        base = self.ensure_schema(base_structured)
        current = self.ensure_schema(current_structured)

        base_skills = {
            canonicalize_skill(skill)
            for category in base["skills"].get("categories", [])
            for skill in category.get("items", [])
            if skill.strip()
        }
        current_skills = {
            canonicalize_skill(skill)
            for category in current["skills"].get("categories", [])
            for skill in category.get("items", [])
            if skill.strip()
        }

        base_bullets = self._collect_section_bullets(base)
        current_bullets = self._collect_section_bullets(current)

        added_bullets = sorted(current_bullets - base_bullets)
        removed_bullets = sorted(base_bullets - current_bullets)
        common_bullets = sorted(current_bullets.intersection(base_bullets))

        return {
            "keywords_added": sorted(current_skills - base_skills),
            "skills_added": sorted(current_skills - base_skills),
            "skills_removed": sorted(base_skills - current_skills),
            "sections_changed": sorted(
                {
                    section
                    for section in ("summary", "skills", "experience", "projects", "education", "certifications", "awards")
                    if base.get(section) != current.get(section)
                }
            ),
            "bullet_delta": len(added_bullets) - len(removed_bullets),
            "added_bullets": added_bullets,
            "removed_bullets": removed_bullets,
            "shared_bullets": common_bullets,
        }

    def tailor_resume(
        self,
        base_structured_resume: dict[str, Any],
        jd_text: str,
        strict_mode: bool = True,
    ) -> dict[str, Any]:
        base_structured = self.ensure_schema(base_structured_resume)
        parsed_jd = self.jd_parser.parse(jd_text, title_hint="Target role")
        profile_data = JobProfileData(
            title=parsed_jd.role_title,
            normalized_title=parsed_jd.normalized_title,
            seniority_level=parsed_jd.seniority_level,
            required_skills=parsed_jd.required_skills,
            optional_skills=parsed_jd.optional_skills,
            responsibilities=parsed_jd.responsibilities,
            years_experience_min=parsed_jd.years_experience_min,
            years_experience_max=parsed_jd.years_experience_max,
        )

        base_plain_text = self.structured_to_plain_text(base_structured)
        scoring = self.scorer.score(base_plain_text, profile_data).model_dump()

        tailored = copy.deepcopy(base_structured)
        tailored["ats_keywords"] = _dedupe(parsed_jd.ats_keywords + tailored.get("ats_keywords", []))
        tailored["experience"]["items"] = [
            {
                **item,
                "bullets": [
                    self._inject_target_keyword(bullet, parsed_jd.required_skills, strict_mode)
                    for bullet in item.get("bullets", [])
                ],
            }
            for item in tailored["experience"].get("items", [])
        ]
        tailored["projects"]["items"] = [
            {
                **item,
                "bullets": [
                    self._inject_target_keyword(bullet, parsed_jd.required_skills, strict_mode)
                    for bullet in item.get("bullets", [])
                ],
            }
            for item in tailored["projects"].get("items", [])
        ]
        tailored["evidence_map"] = self._build_evidence_map(tailored)

        rewrite_suggestions = self._enforce_strict_rewrites(
            suggestions=scoring.get("rewrite_suggestions", []),
            evidence_map=tailored["evidence_map"],
            strict_mode=strict_mode,
        )
        scoring["rewrite_suggestions"] = rewrite_suggestions
        scoring["strict_mode"] = strict_mode
        scoring["role_archetype"] = parsed_jd.role_archetype
        scoring["responsibility_clusters"] = parsed_jd.responsibility_clusters
        scoring["missing_evidence"] = _dedupe(scoring.get("missing_evidence", []))

        if strict_mode:
            scoring["rejection_risks"] = _dedupe(scoring.get("rejection_risks", []))
            scoring["fastest_fixes"] = _dedupe(scoring.get("fastest_fixes", []))

        if self.reasoner.is_enabled():
            self._apply_optional_reasoner(
                scoring=scoring,
                base_plain_text=base_plain_text,
                parsed_jd=parsed_jd,
                strict_mode=strict_mode,
            )

        return {
            "resume_structured_json": tailored,
            "resume_plain_text": self.structured_to_plain_text(tailored),
            "jd_structured_json": {
                "role_title": parsed_jd.role_title,
                "normalized_title": parsed_jd.normalized_title,
                "role_archetype": parsed_jd.role_archetype,
                "seniority_level": parsed_jd.seniority_level,
                "years_experience_required": parsed_jd.years_experience_required,
                "requirements_hard": parsed_jd.requirements_hard,
                "requirements_soft": parsed_jd.requirements_soft,
                "responsibilities": parsed_jd.responsibilities,
                "ats_keywords": parsed_jd.ats_keywords,
                "weight_map": parsed_jd.weight_map,
            },
            "score_snapshot_json": scoring,
            "jd_text_hash": hashlib.sha256(jd_text.encode("utf-8")).hexdigest(),
        }

    def export_docx(self, structured_resume: dict[str, Any], output_path: str) -> None:
        structured = self.ensure_schema(structured_resume)
        doc = Document()
        header = structured["header"]

        if header.get("name"):
            doc.add_heading(header["name"], level=0)
        if header.get("title"):
            doc.add_paragraph(header["title"])
        contact_line = " | ".join(
            [value for value in [header.get("email"), header.get("phone"), header.get("location")] if value]
        )
        if contact_line:
            doc.add_paragraph(contact_line)
        for link in header.get("links", []):
            label = link.get("label", "link")
            url = link.get("url", "")
            if url:
                doc.add_paragraph(f"{label}: {url}")

        if structured.get("summary"):
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(structured["summary"])

        if structured["skills"].get("categories"):
            doc.add_heading("Skills", level=1)
            for category in structured["skills"].get("categories", []):
                doc.add_paragraph(
                    f"{category.get('name', 'Core')}: {', '.join(category.get('items', []))}"
                )

        if structured["experience"].get("items"):
            doc.add_heading("Experience", level=1)
            for item in structured["experience"].get("items", []):
                title = " | ".join(
                    [entry for entry in [item.get("role"), item.get("company"), self._format_dates(item)] if entry]
                )
                if title:
                    doc.add_paragraph(title)
                for bullet in item.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

        if structured["projects"].get("items"):
            doc.add_heading("Projects", level=1)
            for item in structured["projects"].get("items", []):
                if item.get("name"):
                    doc.add_paragraph(item["name"])
                for bullet in item.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

        if structured["education"].get("items"):
            doc.add_heading("Education", level=1)
            for item in structured["education"].get("items", []):
                title = " | ".join(
                    [entry for entry in [item.get("school"), item.get("degree"), self._format_dates(item)] if entry]
                )
                if title:
                    doc.add_paragraph(title)
                for note in item.get("notes", []):
                    doc.add_paragraph(note)

        if structured.get("certifications"):
            doc.add_heading("Certifications", level=1)
            for item in structured["certifications"]:
                doc.add_paragraph(item, style="List Bullet")

        if structured.get("awards"):
            doc.add_heading("Awards", level=1)
            for item in structured["awards"]:
                doc.add_paragraph(item, style="List Bullet")

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))

    def export_pdf(self, structured_resume: dict[str, Any], output_path: str) -> None:
        structured = self.ensure_schema(structured_resume)
        plain_text = self.structured_to_plain_text(structured)
        lines = [line[:115] for line in plain_text.splitlines() if line.strip()]
        self._write_simple_pdf(lines=lines, output_path=output_path)

    def _parse_header(self, lines: list[str], title_hint: str) -> dict[str, Any]:
        header = default_structured_resume(title_hint)["header"]
        if not lines:
            header["name"] = title_hint if title_hint != "Untitled Resume" else ""
            return header

        first_line = lines[0]
        header["name"] = first_line if len(first_line.split()) <= 6 else title_hint
        if len(lines) > 1 and not _EMAIL_RE.search(lines[1]) and not _PHONE_RE.search(lines[1]):
            header["title"] = lines[1]

        for line in lines:
            if not header.get("email"):
                email = _EMAIL_RE.search(line)
                if email:
                    header["email"] = email.group(0)
            if not header.get("phone"):
                phone = _PHONE_RE.search(line)
                if phone:
                    header["phone"] = phone.group(0)
            for url in _URL_RE.findall(line):
                normalized = url if url.startswith("http") else f"https://{url}"
                label = "Link"
                lower = url.lower()
                if "linkedin" in lower:
                    label = "LinkedIn"
                elif "github" in lower:
                    label = "GitHub"
                header["links"].append({"label": label, "url": normalized})
            if not header.get("location"):
                location_candidate = line
                if (
                    "," in location_candidate
                    and not _EMAIL_RE.search(location_candidate)
                    and not _PHONE_RE.search(location_candidate)
                    and not _URL_RE.search(location_candidate)
                ):
                    header["location"] = location_candidate

        header["links"] = _dedupe_links(header["links"])
        return header

    def _parse_skills(self, lines: list[str]) -> list[dict[str, Any]]:
        categories: list[dict[str, Any]] = []
        fallback_items: list[str] = []
        for line in lines:
            clean = _BULLET_RE.sub("", line).strip()
            if ":" in clean and len(clean.split(":", maxsplit=1)[0]) <= 30:
                name, values = clean.split(":", maxsplit=1)
                items = [item.strip() for item in values.split(",") if item.strip()]
                if items:
                    categories.append({"name": name.strip(), "items": _dedupe(items)})
            else:
                fallback_items.extend([item.strip() for item in clean.split(",") if item.strip()])
        if fallback_items:
            categories.append({"name": "Core", "items": _dedupe(fallback_items)})
        return categories

    def _parse_experience(self, lines: list[str]) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        current: dict[str, Any] | None = None

        for raw in lines:
            line = _clean_line(raw)
            if not line:
                continue
            bullet = _BULLET_RE.sub("", line).strip()
            if _BULLET_RE.match(line):
                if current is None:
                    current = {
                        "company": "",
                        "role": "",
                        "location": "",
                        "start": "",
                        "end": "",
                        "bullets": [],
                        "tech": [],
                    }
                    items.append(current)
                current["bullets"].append(bullet)
                continue

            role, company, start, end = self._parse_role_company_dates(line)
            if current is None or current.get("bullets") or role or company:
                current = {
                    "company": company,
                    "role": role,
                    "location": "",
                    "start": start,
                    "end": end,
                    "bullets": [],
                    "tech": [],
                }
                items.append(current)
            else:
                if role and not current.get("role"):
                    current["role"] = role
                if company and not current.get("company"):
                    current["company"] = company

        return [item for item in items if item.get("role") or item.get("company") or item.get("bullets")]

    def _parse_projects(self, lines: list[str]) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        current: dict[str, Any] | None = None
        for raw in lines:
            line = _clean_line(raw)
            if not line:
                continue
            if _BULLET_RE.match(line):
                if current is None:
                    current = {"name": "Project", "link": "", "bullets": [], "tech": []}
                    items.append(current)
                current["bullets"].append(_BULLET_RE.sub("", line).strip())
                continue

            if current is None or current.get("bullets"):
                current = {"name": line, "link": "", "bullets": [], "tech": []}
                items.append(current)
            else:
                current["name"] = f"{current.get('name', '')} {line}".strip()
        return [item for item in items if item.get("name") or item.get("bullets")]

    def _parse_education(self, lines: list[str]) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        for line in lines:
            clean = _BULLET_RE.sub("", line).strip()
            if not clean:
                continue
            degree = ""
            school = clean
            if "," in clean:
                head, tail = clean.split(",", maxsplit=1)
                if any(term in tail.lower() for term in ["bachelor", "master", "phd", "mba", "b.s", "m.s"]):
                    school = head.strip()
                    degree = tail.strip()
            items.append({"school": school, "degree": degree, "start": "", "end": "", "notes": []})
        return items

    def _parse_role_company_dates(self, line: str) -> tuple[str, str, str, str]:
        date_start = ""
        date_end = ""
        date_match = re.search(
            r"((?:19|20)\d{2}|present|current|now)\s*[-–—to]+\s*((?:19|20)\d{2}|present|current|now)",
            line,
            flags=re.IGNORECASE,
        )
        stripped = line
        if date_match:
            date_start = date_match.group(1)
            date_end = date_match.group(2)
            stripped = line.replace(date_match.group(0), "").strip(" |-—")

        role = ""
        company = ""
        if " at " in stripped.lower():
            parts = re.split(r"\bat\b", stripped, maxsplit=1, flags=re.IGNORECASE)
            role = parts[0].strip(" |-—")
            company = parts[1].strip(" |-—")
        elif "|" in stripped:
            left, right = [entry.strip() for entry in stripped.split("|", maxsplit=1)]
            if len(left.split()) <= 8:
                role = left
                company = right
            else:
                company = left
                role = right
        elif " - " in stripped:
            left, right = [entry.strip() for entry in stripped.split(" - ", maxsplit=1)]
            role = left
            company = right
        else:
            role = stripped
        return role, company, date_start, date_end

    def _collect_section_bullets(self, structured: dict[str, Any]) -> set[str]:
        bullets: set[str] = set()
        for item in structured["experience"].get("items", []):
            for bullet in item.get("bullets", []):
                text = bullet.strip().lower()
                if text:
                    bullets.add(text)
        for item in structured["projects"].get("items", []):
            for bullet in item.get("bullets", []):
                text = bullet.strip().lower()
                if text:
                    bullets.add(text)
        return bullets

    def _compute_ats_keywords(self, structured: dict[str, Any]) -> list[str]:
        keywords: list[str] = []
        for category in structured["skills"].get("categories", []):
            keywords.extend(category.get("items", []))
        for item in structured["experience"].get("items", []):
            keywords.extend(item.get("tech", []))
        return _dedupe([canonicalize_skill(keyword) for keyword in keywords if keyword])

    def _build_evidence_map(self, structured: dict[str, Any]) -> dict[str, list[str]]:
        evidence: dict[str, list[str]] = {}
        for exp_index, item in enumerate(structured["experience"].get("items", [])):
            for bullet_index, bullet in enumerate(item.get("bullets", [])):
                key = f"experience.{exp_index}.bullet.{bullet_index}"
                evidence[key] = [bullet]
        for proj_index, item in enumerate(structured["projects"].get("items", [])):
            for bullet_index, bullet in enumerate(item.get("bullets", [])):
                key = f"projects.{proj_index}.bullet.{bullet_index}"
                evidence[key] = [bullet]
        return evidence

    @staticmethod
    def _format_dates(item: dict[str, Any]) -> str:
        start = (item.get("start") or "").strip()
        end = (item.get("end") or "").strip()
        if start and end:
            return f"{start} — {end}"
        return start or end

    @staticmethod
    def _inject_target_keyword(
        bullet: str,
        required_skills: list[str],
        strict_mode: bool,
    ) -> str:
        if not strict_mode:
            return bullet
        lower = bullet.lower()
        for skill in required_skills[:8]:
            canonical = canonicalize_skill(skill)
            if canonical in lower:
                return bullet
            aliases = expand_skill_aliases(canonical)
            if any(alias in lower for alias in aliases if alias != canonical):
                return f"{bullet} ({canonical})"
        return bullet

    @staticmethod
    def _enforce_strict_rewrites(
        suggestions: list[dict[str, Any]],
        evidence_map: dict[str, list[str]],
        strict_mode: bool,
    ) -> list[dict[str, Any]]:
        if not suggestions:
            return []
        evidence_text = " ".join(
            snippet for snippets in evidence_map.values() for snippet in snippets
        ).lower()
        normalized: list[dict[str, Any]] = []
        for entry in suggestions:
            current = dict(entry)
            example = str(current.get("example_bullet", "")).strip()
            requirement = str(current.get("requirement", "")).strip()
            if strict_mode:
                requirement_tokens = [token for token in re.split(r"\W+", requirement.lower()) if len(token) > 2]
                has_grounding = bool(requirement_tokens) and any(token in evidence_text for token in requirement_tokens)
                if not has_grounding:
                    if example:
                        example = re.sub(r"^\[Add if true\]\s*", "", example, flags=re.IGNORECASE)
                        current["example_bullet"] = f"[Add if true] {example}"
                    else:
                        current["example_bullet"] = f"[Add if true] Demonstrate {requirement.lower()} with a specific result."
            normalized.append(current)
        return normalized

    def _apply_optional_reasoner(
        self,
        scoring: dict[str, Any],
        base_plain_text: str,
        parsed_jd: Any,
        strict_mode: bool,
    ) -> None:
        reasoned = self.reasoner.generate(
            {
                "strict_mode": strict_mode,
                "facts_resume": base_plain_text[:3500],
                "facts_jd": {
                    "role_title": parsed_jd.role_title,
                    "required_skills": parsed_jd.required_skills[:30],
                    "responsibilities": parsed_jd.responsibilities[:30],
                },
                "current_summary": scoring.get("explanation_summary", ""),
                "current_verdict": scoring.get("recruiter_verdict", ""),
            }
        )
        if not reasoned:
            return
        if isinstance(reasoned.get("explanation_summary"), str):
            scoring["explanation_summary"] = reasoned["explanation_summary"].strip()
        if isinstance(reasoned.get("recruiter_verdict"), str):
            scoring["recruiter_verdict"] = reasoned["recruiter_verdict"].strip()
        if isinstance(reasoned.get("rewrite_suggestions"), list):
            scoring["rewrite_suggestions"] = self._enforce_strict_rewrites(
                suggestions=[
                    entry for entry in reasoned["rewrite_suggestions"] if isinstance(entry, dict)
                ],
                evidence_map={"resume": [base_plain_text]},
                strict_mode=strict_mode,
            )

    @staticmethod
    def _write_simple_pdf(lines: list[str], output_path: str) -> None:
        def escape_pdf_text(value: str) -> str:
            return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        content_lines = ["BT", "/F1 10 Tf", "14 TL", "50 780 Td"]
        max_lines = min(len(lines), 260)
        for idx in range(max_lines):
            text = escape_pdf_text(lines[idx])
            if idx > 0:
                content_lines.append("T*")
            content_lines.append(f"({text}) Tj")
        content_lines.append("ET")

        content = "\n".join(content_lines).encode("latin-1", errors="replace")
        objects: list[bytes] = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
            b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
        ]

        output = io.BytesIO()
        output.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets: list[int] = []

        for obj_id, obj in enumerate(objects, start=1):
            offsets.append(output.tell())
            output.write(f"{obj_id} 0 obj\n".encode("ascii"))
            output.write(obj)
            output.write(b"\nendobj\n")

        xref_position = output.tell()
        output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        output.write(b"0000000000 65535 f \n")
        for offset in offsets:
            output.write(f"{offset:010d} 00000 n \n".encode("ascii"))

        output.write(
            (
                "trailer\n"
                f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_position}\n%%EOF\n"
            ).encode("ascii")
        )

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(output.getvalue())


def _dedupe_links(links: list[dict[str, str]]) -> list[dict[str, str]]:
    seen: set[str] = set()
    output: list[dict[str, str]] = []
    for link in links:
        url = link.get("url", "").strip()
        if not url or url in seen:
            continue
        seen.add(url)
        output.append({"label": link.get("label", "Link"), "url": url})
    return output
